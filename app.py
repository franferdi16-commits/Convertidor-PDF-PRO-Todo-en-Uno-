import streamlit as st
import pdfplumber
import pytesseract
import io
import fitz  # PyMuPDF
import os
from PIL import Image
from docx import Document

# --- 1. CONFIGURACIÓN DEL MOTOR OCR ---
posibles_rutas = [
    r"C:/Program Files/Tesseract-OCR/tesseract.exe",
    r"C:/Tesseract-OCR/tesseract.exe",
]
pytesseract.pytesseract.tesseract_cmd = ""
for ruta in posibles_rutas:
    if os.path.exists(ruta):
        pytesseract.pytesseract.tesseract_cmd = ruta
        break

# --- 2. CUERPO PRINCIPAL DE LA APP ---
st.title("🚀 Convertidor PDF PRO (Todo-en-Uno)")
st.markdown("---")
st.image("https://portalguarani.com/userfiles/images/Norma%20Buttner/05%20Norma%20Buttner%20Lapacho%20amarillo%2070%20x%50%20cm%20portalguarani.jpg", caption="Óleo de Norma Büttner - Lapacho Amarillo")

pdfs = st.file_uploader("Sube tus archivos PDF aquí", type="pdf", accept_multiple_files=True)

if pdfs:
    for pdf in pdfs:
        with st.expander(f"📦 Procesando: {pdf.name}", expanded=True):
            n_limpio = "".join(c for c in pdf.name if c.isalnum() or c in "._-").replace(".pdf", "")

            doc_word = Document()
            texto_respaldo = []

            # Leer PDF directo en memoria
            pdf_bytes = io.BytesIO(pdf.getvalue())

            try:
                # PASO A: TEXTO Y TABLAS
                with st.spinner('Extrayendo texto y tablas nativas...'):
                    with pdfplumber.open(pdf_bytes) as pb:
                        for page in pb.pages:
                            tables = page.extract_tables()
                            for table in tables:
                                if table:
                                    t_word = doc_word.add_table(rows=len(table), cols=len(table[0]) if table[0] else 0)
                                    t_word.style = 'Table Grid'
                                    for i, row in enumerate(table):
                                        for j, cell in enumerate(row):
                                            t_word.cell(i, j).text = str(cell) if cell else ""

                            p_text = page.extract_text()
                            if p_text:
                                doc_word.add_paragraph(p_text)
                                texto_respaldo.append(p_text)

                # PASO B: OCR DE IMÁGENES
                with st.spinner('Escaneando contenido de imágenes (OCR)...'):
                    with fitz.open(stream=pdf_bytes.getvalue()) as doc_fitz:
                        for page in doc_fitz:
                            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                            img = Image.open(io.BytesIO(pix.tobytes("png")))
                            ocr_res = pytesseract.image_to_string(img, lang='spa')
                            if ocr_res.strip():
                                doc_word.add_paragraph(f"[Contenido Escaneado]:\n{ocr_res}")
                                texto_respaldo.append(ocr_res)

                # PASO C: GENERACIÓN DE DESCARGAS
                st.success(f"¡{pdf.name} procesado correctamente!")
                word_io = io.BytesIO()
                doc_word.save(word_io)

                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "📥 Descargar Word (.docx)",
                        word_io.getvalue(),
                        f"{n_limpio}_EDITABLE.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with col2:
                    st.download_button(
                        "📄 Descargar Texto (.txt)",
                        "\n".join(texto_respaldo).encode("utf-8"),
                        f"{n_limpio}_RESPALDO.txt"
                    )
            except Exception as e:
                st.error(f"Error procesando {pdf.name}: {e}")

# Copyright
st.markdown("---")
st.caption("© Hecho en Paraguay por Alberto Fernández 🌟")
